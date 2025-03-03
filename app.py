__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
import streamlit as st
import os
import json
import markdown
import base64
from crewai import LLM, Agent, Crew, Process, Task
from crewai_tools import RagTool
from dotenv import load_dotenv
from templates import query_template, summarization_template
import io
from docx import Document
import uuid
import time as import_time

load_dotenv()

# Generate a unique session ID for user if not exists
if 'session_id' not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# Initialize session state
if 'documents' not in st.session_state:
    st.session_state.documents = []
if 'processing' not in st.session_state:
    st.session_state.processing = False
if 'results' not in st.session_state:
    st.session_state.results = None
if 'api_key_set' not in st.session_state:
    st.session_state.api_key_set = False
if 'history' not in st.session_state:
    st.session_state.history = []

# Page configuration
st.set_page_config(
    page_title="InfoLegal Agent",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for a cleaner, modern design
st.markdown("""
<style>
    /* Modern clean theme with legal blue accent */
    .stApp {
        background-color: #f5f7fa;
        color: #333;
    }
    
    /* Header styling */
    .app-header {
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 1rem 0;
        border-bottom: 1px solid #eaeaea;
        margin-bottom: 2rem;
    }
    
    .app-header h1 {
        color: #1e3a8a;
        margin: 0;
        font-weight: 600;
        font-size: 1.8rem;
    }
    
    .app-logo {
        font-size: 2.5rem;
        color: #1e3a8a;
    }
    
    /* Card styling */
    .content-card {
        background-color: white;
        border-radius: 8px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        padding: 1.5rem;
        margin-bottom: 1.5rem;
    }
    
    /* Button styling */
    .stButton > button {
        background-color: #1e3a8a;
        color: white;
        border-radius: 6px;
        border: none;
        padding: 0.5rem 1rem;
        font-weight: 500;
    }
    
    .stButton > button:hover {
        background-color: #1e40af;
        box-shadow: 0 2px 10px rgba(30, 58, 138, 0.2);
    }
    
    .secondary-button > button {
        background-color: #e2e8f0;
        color: #1e3a8a;
    }
    
    .secondary-button > button:hover {
        background-color: #cbd5e1;
    }
    
    /* Input fields */
    .stTextInput > div > div > input, .stTextArea > div > div > textarea {
        border-radius: 6px;
        border: 1px solid #e2e8f0;
    }
    
    /* Sidebar styling */
    section[data-testid="stSidebar"] {
        background-color: white;
        border-right: 1px solid #eaeaea;
    }
    
    /* History item styling */
    .history-item {
        padding: 0.75rem;
        border-radius: 6px;
        background-color: #f1f5f9;
        margin-bottom: 0.5rem;
        cursor: pointer;
        transition: background-color 0.2s;
    }
    
    .history-item:hover {
        background-color: #e2e8f0;
    }
    
    /* Results styling */
    .result-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 1rem;
    }
    
    .result-content {
        background-color: #f8fafc;
        border-radius: 6px;
        padding: 1rem;
        border-left: 4px solid #1e3a8a;
    }
    
    /* Document list styling */
    .doc-list-item {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding: 0.5rem;
        border-radius: 4px;
        background-color: #f1f5f9;
        margin-bottom: 0.5rem;
    }
    
    /* Footer styling */
    .footer {
        text-align: center;
        padding-top: 2rem;
        color: #64748b;
        font-size: 0.875rem;
        border-top: 1px solid #eaeaea;
        margin-top: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# Logo and title
st.markdown("""
<div class="legal-header">
    <img src="https://w7.pngwing.com/pngs/254/724/png-transparent-lawyer-statute-legislature-procedural-law-lawyers-silhouette-people-logo-man-silhouette.png" alt="InfoLegal Logo">
    <h1>InfoLegal Agent</h1>
</div>
""", unsafe_allow_html=True)

# Sidebar with tabs
with st.sidebar:
    st.markdown("### InfoLegal Agent")
    
    tab1, tab2 = st.tabs(["Documents", "History"])
    
    with tab1:
        # Document upload section
        st.markdown("#### Add Documents")
        
        doc_source = st.radio(
            "Source",
            ["Upload File", "Web URL"],
            horizontal=True
        )
        
        if doc_source == "Upload File":
            uploaded_file = st.file_uploader(
                "Upload document", 
                type=["pdf", "csv", "json", "txt", "docx"],
                label_visibility="collapsed"
            )
            
            if uploaded_file is not None:
                # Save the file temporarily
                file_path = f"temp_{uploaded_file.name}"
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                    
                add_file = st.button("Add Document", use_container_width=True)
                if add_file:
                    file_type = uploaded_file.name.split('.')[-1].lower()
                    st.session_state.documents.append({
                        "type": "file",
                        "path": file_path,
                        "name": uploaded_file.name,
                        "file_type": file_type
                    })
                    st.success(f"Added: {uploaded_file.name}")
        
        else:  # Web URL
            url_input = st.text_input(
                "Website URL",
                placeholder="https://example.com"
            )
            
            add_url = st.button("Add URL", use_container_width=True)
            if add_url and url_input:
                if url_input.startswith(("http://", "https://")):
                    st.session_state.documents.append({
                        "type": "web_page",
                        "url": url_input,
                        "name": url_input.split("/")[-1] if url_input.split("/")[-1] else url_input
                    })
                    st.success(f"Added: {url_input}")
                else:
                    st.error("Please enter a valid URL")
        
        # Current documents section
        if st.session_state.documents:
            st.markdown("#### Current Documents")
            
            for i, doc in enumerate(st.session_state.documents):
                st.markdown(f"""
                <div class="doc-list-item">
                    <span>{'📄' if doc['type'] == 'file' else '🌐'} {doc['name']}</span>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns([1, 1])
                with col1:
                    if st.button("Remove", key=f"remove_{i}", use_container_width=True):
                        st.session_state.documents.pop(i)
                        st.experimental_rerun()
            
            if st.button("Clear All", use_container_width=True):
                st.session_state.documents = []
                st.experimental_rerun()
    
    with tab2:
        st.markdown("#### Query History")
        
        if not st.session_state.history:
            st.info("No queries yet. Your history will appear here.")
        else:
            for i, item in enumerate(st.session_state.history):
                if st.button(
                    f"Query: {item['query'][:30]}{'...' if len(item['query']) > 30 else ''}",
                    key=f"history_{i}",
                    use_container_width=True
                ):
                    # Set the query text and results from history
                    st.session_state.query_text = item['query']
                    st.session_state.results = item['results']
                    st.experimental_rerun()

# Main content
col1, col2 = st.columns([2, 1])

with col1:
    # Query input card
    st.markdown('<div class="content-card">', unsafe_allow_html=True)
    st.markdown("### Ask Your Legal Question")
    
    # Check if we have a query from history
    if 'query_text' in st.session_state:
        query_text = st.session_state.query_text
        # Clear the temporary variable
        del st.session_state.query_text
    else:
        query_text = ""
    
    user_query = st.text_area(
        "Enter your inquiry",
        value=query_text,
        placeholder="E.g., What are the steps to file a civil lawsuit in India?",
        height=120,
        label_visibility="collapsed"
    )
    
    col_a, col_b, col_c = st.columns([1, 1, 2])
    
    with col_a:
        process_button = st.button(
            "Submit Query", 
            disabled=not st.session_state.documents or not user_query,
            use_container_width=True
        )
    
    with col_b:
        st.markdown('<div class="secondary-button">', unsafe_allow_html=True)
        clear_button = st.button("Clear", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        if clear_button:
            user_query = ""
            st.session_state.results = None
            st.experimental_rerun()
    
    if not st.session_state.documents:
        st.warning("⚠️ Please add at least one document in the sidebar.")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Results section
    if process_button and user_query:
        st.session_state.processing = True
    
    if st.session_state.processing:
        st.markdown('<div class="content-card">', unsafe_allow_html=True)
        st.markdown("### Processing Results")
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Update progress
            status_text.text("Setting up tools and agents...")
            progress_bar.progress(10)
            
            # Initialize tools and agents
            config = {
                "llm": {
                    "provider": "google",
                    "config": {
                        "model": "gemini-1.5-pro",
                        "api_key": os.environ.get("GEMINI_API_KEY"), 
                    }
                },
                "embedder": {
                    "provider": "huggingface",
                    "config": {
                        "model": "sentence-transformers/all-MiniLM-L6-v2",
                    }
                }
            }
            
            rag_tool = RagTool(config=config)
            
            # Update progress
            status_text.text("Processing documents...")
            progress_bar.progress(30)
            
            # Add documents to RAG tool - FIXED with source parameter
            for doc in st.session_state.documents:
                if doc["type"] == "file":
                    # Determine the source type based on file extension
                    file_type = doc.get("file_type", "").lower()
                    if file_type == "pdf":
                        source_type = "pdf_file"
                    elif file_type == "docx":
                        source_type = "docx"
                    elif file_type == "txt":
                        source_type = "text"
                    elif file_type == "csv":
                        source_type = "csv"
                    elif file_type == "json":
                        source_type = "json"
                    elif file_type == "mdx":
                        source_type = "mdx"
                    else:
                        source_type = "document"  # default source type
                        
                    rag_tool.add(source=source_type, data_type="file", path=doc["path"])
                elif doc["type"] == "web_page":
                     rag_tool.add(doc["url"],data_type="web_page")
            
            # Update progress
            status_text.text("Analyzing your query...")
            progress_bar.progress(50)
            
            # Initialize LLM
            llm = LLM(model="google/gemini-1.5-pro", api_key=os.environ.get("GEMINI_API_KEY"), temperature=0.25)
            
            # Initialize agents
            query_agent = Agent(
                role="Legal Information Retrieval Specialist - India Focus",
                goal="Efficiently and accurately retrieve relevant sections from Indian legal documents",
                backstory="You are a highly specialized AI agent with deep indexing and search capabilities within a specific corpus of Indian legal texts."
                "You understand legal terminology and can quickly pinpoint the most relevant passages, sections, and clauses related to a user's natural language query."
                "You prioritize precision and recall, ensuring that all potentially relevant information is identified" 
                "while minimizing irrelevant results. You are optimized for speed and accuracy in a legal context..",
                llm=llm,
                memory=True,  
                respect_context_window=True, 
                use_system_prompt=True,
                cache=True, 
                system_template=query_template.system_template,  
                prompt_template=query_template.prompt_template,  
                response_template=query_template.response_template,
                tools=[rag_tool]
            )
            
            # Update progress
            status_text.text("Finding relevant legal information...")
            progress_bar.progress(70)
            
            summarization_agent = Agent(
                role="Legal Text Simplification and Explanation Expert",
                goal="Transform complex legal passages into clear, concise, and easy-to-understand summaries, preserving accuracy.",
                backstory="You are an advanced AI agent trained in legal text summarization and simplification. You understand legal principles and can explain complex concepts in plain language."
                "You prioritize making legal information accessible to non-experts without losing crucial details or introducing inaccuracies."
                "You are adept at handling legal jargon and converting it into understandable terms. You can structure summaries logically.",
                llm=llm, 
                memory=True, 
                respect_context_window=True, 
                use_system_prompt=True,
                cache=True,  
                system_template=summarization_template.system_template, 
                prompt_template=summarization_template.prompt_template,
                response_template=summarization_template.response_template,  
                tools=[rag_tool],
                multimodal=True
            )
            
            # Initialize tasks
            query_task = Task(
                description="""
                    Retrieve relevant sections from Indian legal documents based on the user's query.
                    Focus on accuracy and completeness. Find ALL relevant sections, passages, and clauses.
                    Do NOT summarize the content. Only extract the verbatim text and provide citations.
                """,
                expected_output="""
                    A JSON object containing the original user query and a list of results.  
                    Each result should include the document title, a full citation (section, chapter, page, etc.),
                    and the exact text of the relevant section.  Follow this format:

                    {
                      "query": "<User's original query>",
                      "results": [
                        {
                          "document": "<Document Title>",
                          "citation": "<Full Citation (Section, Chapter, Page, etc.)>",
                          "extract": "<Exact text of the relevant section>"
                        },
                        {
                          "document": "<Document Title>",
                          "citation": "<Full Citation (Section, Chapter, Page, etc.)>",
                          "extract": "<Exact text of the relevant section>"
                        }
                      ]
                    }
                """,
                agent=query_agent,
            )
            
            # Update progress
            status_text.text("Generating user-friendly summary...")
            progress_bar.progress(90)
            
            summarization_task = Task(
                description="""
                    Take the output from the Legal Information Retrieval Specialist (Query Agent) and 
                    transform it into a clear, concise, and easy-to-understand summary.  
                    Preserve the original legal meaning and accuracy.  
                    Replace legal jargon with plain language where possible.
                    Structure the summary logically (e.g., using bullet points or numbered steps if appropriate).
                    Conclude by offering to provide more details.
                """,
                expected_output="""
                    A plain text summary of the legal information, followed by a question prompting the user
                    if they want more details. For example:

                    Filing a lawsuit in India involves preparing legal documents, submitting a petition in court,
                    serving a notice to the opposing party, and attending hearings.

                    Would you like more details on any step?
                """,
                agent=summarization_agent,
                context=[query_task],
                output_file="summarization.md"
            )
            
            # Initialize crew
            crew = Crew(
                agents=[query_agent, summarization_agent],
                tasks=[query_task, summarization_task],
                process=Process.sequential,
                ktools=[rag_tool],
                memory=True,
                cache=True
            )
            
            # Run the query
            result = crew.kickoff(inputs=user_query)
            clean_result = markdown.markdown(result)
            
            # Store the result
            st.session_state.results = {
                "raw": result,
                "html": clean_result
            }
            
            # Add to history
            st.session_state.history.append({
                "query": user_query,
                "results": st.session_state.results,
                "timestamp": import_time.time()
            })
            
            # Update progress
            progress_bar.progress(100)
            status_text.text("Done!")
            
            # Reset processing state
            st.session_state.processing = False
            
            # Force a rerun to show results
            st.experimental_rerun()
            
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            st.session_state.processing = False
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Display results
    if st.session_state.results:
        st.markdown('<div class="content-card">', unsafe_allow_html=True)
        
        st.markdown('<div class="result-header">', unsafe_allow_html=True)
        st.markdown("### Results")
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown('<div class="result-content">', unsafe_allow_html=True)
        st.markdown(st.session_state.results["html"], unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        
        col_a, col_b, col_c = st.columns([1, 1, 2])
        
        with col_a:
            # Export as DOCX
            if st.button("Download as DOCX", use_container_width=True):
                # Create a DOCX file
                doc = Document()
                doc.add_heading('Legal Query Results', 0)
                
                # Add query
                doc.add_heading('Query', level=1)
                doc.add_paragraph(user_query)
                
                # Add results (stripping HTML)
                doc.add_heading('Response', level=1)
                doc.add_paragraph(st.session_state.results["raw"])
                
                # Save to BytesIO
                docx_file = io.BytesIO()
                doc.save(docx_file)
                docx_file.seek(0)
                
                # Create download button
                st.download_button(
                    label="Download DOCX",
                    data=docx_file,
                    file_name="legal_results.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with col_b:
            # Export as Markdown
            if st.button("Download as Markdown", use_container_width=True):
                st.download_button(
                    label="Download Markdown",
                    data=st.session_state.results["raw"],
                    file_name="legal_results.md",
                    mime="text/markdown",
                    use_container_width=True
                )
                
        st.markdown('</div>', unsafe_allow_html=True)

with col2:
    # Tips and information card
    st.markdown('<div class="content-card">', unsafe_allow_html=True)
    st.markdown("### How to Use InfoLegal")
    
    st.markdown("""
    #### Quick Guide
    
    1. Add legal documents via upload or URL
    2. Enter your legal query
    3. Review the simplified legal information
    4. Download results if needed
    
    #### Tips for Best Results
    
    • Be specific in your legal questions
    • Add relevant legal documents
    • For India-specific questions, include Indian legal texts
    • Browse your query history to revisit past queries
    """)
    
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("""
<div class="footer">
    <p>© 2025 InfoLegal Agent | Powered by Qest.ai</p>
</div>
""", unsafe_allow_html=True)
